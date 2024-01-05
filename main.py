import os
import datetime
import pandas as pd
from tkinter import *
from tkinter import filedialog, messagebox
from tkinter import ttk
from tkcalendar import Calendar  # Импортируем библиотеку для выбора даты
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import threading


def browse_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    input_path_entry.delete(0, END)
    input_path_entry.insert(0, file_path)


def browse_word_template():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    template_path_entry.delete(0, END)
    template_path_entry.insert(0, file_path)


def generate_word_files():
    input_excel_file = input_path_entry.get()
    if not input_excel_file:
        messagebox.showerror("Ошибка", "Выберите файл Excel!")
        return

    template_word_file = template_path_entry.get()
    if not template_word_file:
        messagebox.showerror("Ошибка", "Выберите файл-шаблон Word!")
        return

    current_date = datetime.datetime.now().strftime('%d.%m.%y')
    output_folder = os.path.join(os.getcwd(), current_date)
    os.makedirs(output_folder, exist_ok=True)

    df = pd.read_excel(input_excel_file)

    for index, row in df.iterrows():
        doc = Document(template_word_file)

        for para in doc.paragraphs:
            para.text = para.text.replace("{Account_Number}", str(row["Лиц. счет"]))
            para.text = para.text.replace("{Address}", str(row["Адрес"]))
            para.text = para.text.replace("{Debt_Amount}", str(row["Сумма долга"]))

            # Заменяем плейсхолдер {Status_date} на выбранную дату
            if "{Status_date}" in para.text:
                status_date_str = status_date_calendar.get_date()
                para.text = para.text.replace("{Status_date}", status_date_str)

        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        for para in doc.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        account_number = row["Лиц. счет"]
        word_filename = os.path.join(output_folder, f'Повідомлення - {account_number}.docx')
        doc.save(word_filename)

    messagebox.showinfo("Готово", "Word файлы сохранены!")


def start_processing_thread():
    processing_thread = threading.Thread(target=generate_word_files)
    processing_thread.start()


# Создание окна
root = Tk()
root.title("Генератор Word файлов")
root.geometry("800x800")  # Устанавливаем размер окна

# Создание метки и поля для выбора файла Excel
input_label = Label(root, text="Выберите файл Excel:")
input_label.pack(pady=(20, 5))  # Устанавливаем вертикальное отступы

input_path_entry = Entry(root, width=50)
input_path_entry.pack()

browse_excel_button = ttk.Button(root, text="Обзор", command=browse_excel_file)
browse_excel_button.pack(pady=(5, 10))

# Создание метки и поля для выбора файла-шаблона Word
template_label = Label(root, text="Выберите файл-шаблон Word:")
template_label.pack()

template_path_entry = Entry(root, width=50)
template_path_entry.pack()

browse_template_button = ttk.Button(root, text="Обзор", command=browse_word_template)
browse_template_button.pack(pady=(5, 10))

# Выбор даты
status_date_label = Label(root, text="Выберите дату:")
status_date_label.pack()

status_date_calendar = Calendar(root)
status_date_calendar.pack()

# Кнопка для генерации Word файлов
generate_button = ttk.Button(root, text="Создать Word файлы", command=start_processing_thread)
generate_button.pack()

root.mainloop()
